#!/usr/bin/env python3
"""
Компактный демонстрационный документ для научной статьи.
Все элементы на одной странице (~1-2 стр).
Содержит триггеры для всех процессоров модуля.

Использование:
    .venv/bin/python create_compact_demo.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def create_compact_demo():
    """Создать компактный демонстрационный документ."""
    doc = Document()
    
    # ========== МИНИМАЛЬНЫЙ ТИТУЛ ==========
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run('СТАНДАРТ РЕДАКТИРОВАНИЯ')
    title_run.bold = True
    title_run.font.size = Pt(14)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.add_run('Демонстрация автоматического форматирования')
    
    doc.add_paragraph()  # Пустая строка
    
    # ========== ОСНОВНОЕ СОДЕРЖАНИЕ ==========
    
    # Раздел 1 (Heading 1) - триггер для SectionProcessor
    h1 = doc.add_heading('Введение', level=1)
    doc.add_paragraph(
        'Данный документ демонстрирует полный функционал системы автоматического '
        'форматирования документов Word согласно ГОСТ.'
    )
    
    # Подраздел 1.1 (Heading 2)
    h2 = doc.add_heading('Цель работы', level=2)
    doc.add_paragraph(
        'Показать применение многоуровневой нумерации разделов, '
        'автоматического оглавления и прочих функций.'
    )
    
    # Подподраздел 1.1.1 (Heading 3)
    h3 = doc.add_heading('Основные компоненты', level=3)
    doc.add_paragraph('StyleProcessor — форматирование стилей', style='List Bullet')
    doc.add_paragraph('SectionProcessor — нумерация разделов', style='List Bullet')
    doc.add_paragraph('TOCProcessor — оглавление', style='List Bullet')
    
    # Раздел 2 (Heading 1)
    doc.add_heading('Методы', level=1)
    doc.add_paragraph(
        'Система использует многомодульную архитектуру для обработки документов.'
    )
    
    # Подраздел 2.1 (Heading 2)
    doc.add_heading('Обработка данных', level=2)
    
    # Таблица - триггер для StyleProcessor
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Заголовок таблицы
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Параметр'
    hdr_cells[1].text = 'Значение'
    
    # Данные таблицы
    data = [
        ('Версия', '2.0'),
        ('Статус', 'Активно'),
        ('Уровни', '3'),
    ]
    
    for i, (param, value) in enumerate(data, 1):
        row = table.rows[i].cells
        row[0].text = param
        row[1].text = value
    
    doc.add_paragraph()  # Пустая строка
    
    # Раздел 3 (Heading 1)
    doc.add_heading('Результаты', level=1)
    doc.add_paragraph(
        'После обработки документ получает правильное форматирование, '
        'нумерацию и структуру в соответствии с конфигурацией.'
    )
    
    doc.add_paragraph()  # Пустая строка
    
    # ========== ПРИЛОЖЕНИЯ ==========
    # Триггеры для AppendixProcessor (ключевое слово "Приложение")
    
    doc.add_heading('Приложение А', level=1)
    doc.add_paragraph('Дополнительные сведения и рекомендации.')
    
    doc.add_heading('Приложение Б', level=1)
    doc.add_paragraph('Справочная информация и примеры использования.')
    
    # Сохранить
    output_path = Path('doc_editor/tests/test_data/demo_document.docx')
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    
    print(f"✅ Компактный демонстрационный документ создан:")
    print(f"   {output_path.absolute()}")
    print(f"\n📋 Содержит триггеры для всех процессоров:")
    print(f"   ✓ Heading 1, 2, 3 → SectionProcessor (многоуровневая нумерация)")
    print(f"   ✓ Таблица → StyleProcessor (форматирование)")
    print(f"   ✓ Заголовки → TOCProcessor (оглавление)")
    print(f"   ✓ 'Приложение А/Б' → AppendixProcessor (нумерация приложений)")


if __name__ == '__main__':
    create_compact_demo()
