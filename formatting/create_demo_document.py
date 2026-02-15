#!/usr/bin/env python3
"""
Скрипт для создания демонстрационного документа Word для тестирования модуля.

Использование:
    python create_demo_document.py [output_path] [--minimal]

Примеры:
    python create_demo_document.py doc_editor/tests/test_data/demo_document.docx
    python create_demo_document.py doc_editor/tests/test_data/demo_document.docx --minimal
"""

import sys
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_heading_with_style(doc: Document, text: str, level: int) -> None:
    """Добавить заголовок с нужным уровнем."""
    if level == 0:
        doc.add_heading(text, level=1)  # Heading 1
    elif level == 1:
        doc.add_heading(text, level=2)  # Heading 2
    else:
        doc.add_heading(text, level=3)  # Heading 3


def add_table_example(doc: Document) -> None:
    """Добавить пример таблицы в документ."""
    doc.add_paragraph()
    doc.add_paragraph("Таблица 1 - Пример данных").bold = True
    
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Заголовок таблицы
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Параметр'
    header_cells[1].text = 'Значение'
    header_cells[2].text = 'Единица'
    
    # Данные таблицы
    data = [
        ('Версия', '1.0', 'текст'),
        ('Дата', '2026-02-13', 'дата'),
        ('Статус', 'Активно', 'текст'),
        ('Уровень', '3', 'число'),
    ]
    
    for i, (param, value, unit) in enumerate(data, 1):
        cells = table.rows[i].cells
        cells[0].text = param
        cells[1].text = value
        cells[2].text = unit


def create_minimal_demo_document(output_path: str) -> None:
    """Создать минимальный демонстрационный документ."""
    print(f"📄 Создаю минимальный демонстрационный документ...")
    
    doc = Document()
    
    # Титул
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run('ДЕМОНСТРАЦИОННЫЙ ДОКУМЕНТ')
    title_run.bold = True
    title_run.font.size = Pt(16)
    
    doc.add_paragraph()  # Пустая строка
    
    # Раздел 1
    add_heading_with_style(doc, 'Раздел 1. Введение', 0)
    doc.add_paragraph('Это демонстрационный раздел документа. Содержит текст для тестирования форматирования и нумерации.')
    
    # Подраздел 1.1
    add_heading_with_style(doc, 'Подраздел 1.1. Цель', 1)
    doc.add_paragraph('Цель этого документа - продемонстрировать все функции модуля редактирования.')
    
    # Раздел 2
    add_heading_with_style(doc, 'Раздел 2. Основное содержание', 0)
    doc.add_paragraph('Второй раздел основного документа.')
    
    # Подраздел 2.1
    add_heading_with_style(doc, 'Подраздел 2.1. Данные', 1)
    doc.add_paragraph('Здесь размещены важные данные.')
    
    # Таблица
    add_table_example(doc)
    
    # Раздел 3
    add_heading_with_style(doc, 'Раздел 3. Завершение', 0)
    doc.add_paragraph('Третий раздел документа.')
    
    doc.add_paragraph()  # Пустая строка
    
    # Приложения
    add_heading_with_style(doc, 'Приложение А. Дополнительная информация', 0)
    doc.add_paragraph('Содержимое первого приложения.')
    
    add_heading_with_style(doc, 'Приложение Б. Справочные данные', 0)
    doc.add_paragraph('Содержимое второго приложения.')
    
    doc.save(output_path)
    print(f"✅ Минимальный документ создан: {output_path}")


def create_full_demo_document(output_path: str) -> None:
    """Создать полный демонстрационный документ со всеми функциями."""
    print(f"📄 Создаю полный демонстрационный документ...")
    
    doc = Document()
    
    # ========== ТИТУЛ ==========
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run('ДЕМОНСТРАЦИОННЫЙ ДОКУМЕНТ')
    title_run.bold = True
    title_run.font.size = Pt(18)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle.add_run('Система форматирования документов')
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(14)
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_para.add_run('Версия 1.0 / Февраль 2026')
    
    # Разделитель
    for _ in range(3):
        doc.add_paragraph()
    
    # ========== ОСНОВНОЕ СОДЕРЖАНИЕ ==========
    
    # Раздел 1: Введение
    add_heading_with_style(doc, 'Введение', 0)
    doc.add_paragraph(
        'Данный документ демонстрирует все возможности модуля редактирования документов Word. '
        'Включает примеры нумерации разделов, таблиц, приложений и прочих элементов.'
    )
    
    # Подраздел 1.1
    add_heading_with_style(doc, 'Цель и область применения', 1)
    doc.add_paragraph(
        'Целью данного документа является полная демонстрация функционала системы. '
        'Область применения охватывает все процессоры обработки документов.'
    )
    
    # Подподраздел 1.1.1
    add_heading_with_style(doc, 'Основные характеристики', 2)
    doc.add_paragraph('Система поддерживает:')
    for feature in [
        'Автоматическую многоуровневую нумерацию разделов',
        'Генерацию оглавления на основе заголовков',
        'Добавление предисловия',
        'Автоматическое форматирование приложений',
        'Управление полями и шрифтами документа'
    ]:
        doc.add_paragraph(feature, style='List Bullet')
    
    # Подраздел 1.2
    add_heading_with_style(doc, 'Требования и ограничения', 1)
    doc.add_paragraph('Минимальные требования:')
    doc.add_paragraph('Python 3.8+', style='List Bullet')
    doc.add_paragraph('python-docx', style='List Bullet')
    doc.add_paragraph('PyYAML', style='List Bullet')
    
    # Раздел 2: Нормативные ссылки
    add_heading_with_style(doc, 'Нормативные ссылки', 0)
    doc.add_paragraph('Этот раздел содержит ссылки на используемые стандарты.')
    
    add_heading_with_style(doc, 'Межгосударственные стандарты', 1)
    doc.add_paragraph('ГОСТ Р 1.5-2004 - Межгосударственная система по стандартизации', style='List Bullet')
    
    add_heading_with_style(doc, 'Национальные стандарты', 1)
    doc.add_paragraph('ГОСТ Р 2.103-2013 - Единая система конструкторской документации', style='List Bullet')
    
    # Раздел 3: Термины и определения
    add_heading_with_style(doc, 'Термины и определения', 0)
    doc.add_paragraph('Определения используемых в документе терминов.')
    
    add_heading_with_style(doc, 'Основные термины', 1)
    doc.add_paragraph('Документ - совокупность данных и информации в электронном виде.', style='List Bullet')
    doc.add_paragraph('Процессор - компонент системы для обработки определенного типа элементов.', style='List Bullet')
    
    # Раздел 4: Требования
    add_heading_with_style(doc, 'Требования', 0)
    doc.add_paragraph('Основные требования к обработке документов.')
    
    add_heading_with_style(doc, 'Функциональные требования', 1)
    doc.add_paragraph('Система должна поддерживать все форматы документов Word.')
    doc.add_paragraph('Оглавление должно обновляться автоматически.')
    doc.add_paragraph('Нумерация разделов должна быть многоуровневой.')
    
    add_heading_with_style(doc, 'Общие требования', 2)
    doc.add_paragraph('Все требования безопасности должны быть соблюдены.')
    
    add_heading_with_style(doc, 'Производительность', 2)
    doc.add_paragraph('Время обработки документа не должно превышать 5 секунд.')
    
    add_heading_with_style(doc, 'Специальные требования', 1)
    doc.add_paragraph('Поддержка русского языка обязательна.')
    doc.add_paragraph('Поддержка кириллицы в заголовках обязательна.')
    
    # Раздел 5: Примеры
    add_heading_with_style(doc, 'Примеры использования', 0)
    doc.add_paragraph('Практические примеры применения системы.')
    
    # Таблица примеров
    doc.add_paragraph('Таблица 1 - Примеры конфигурации', style='Heading 2')
    add_table_example(doc)
    
    # Раздел 6
    add_heading_with_style(doc, 'Технические детали', 0)
    doc.add_paragraph('Описание технической реализации.')
    
    add_heading_with_style(doc, 'Архитектура системы', 1)
    doc.add_paragraph('Система построена модульно с использованием паттерна Pipeline.')
    
    add_heading_with_style(doc, 'Компоненты', 2)
    for component in [
        'TitleProcessor - обработка титульного листа',
        'SectionProcessor - нумерация разделов',
        'TOCProcessor - генерация оглавления',
        'PrefaceProcessor - добавление предисловия',
        'AppendixProcessor - обработка приложений',
        'StyleProcessor - форматирование стилей'
    ]:
        doc.add_paragraph(component, style='List Bullet')
    
    # Пустые строки перед приложениями
    for _ in range(2):
        doc.add_paragraph()
    
    # ========== ПРИЛОЖЕНИЯ ==========
    
    add_heading_with_style(doc, 'Приложение А. Содержание конфигурации', 0)
    doc.add_paragraph('Подробное описание структуры конфигурационного файла.')
    doc.add_paragraph('Конфигурация содержит разделы для:')
    for section in ['General Settings', 'Structure', 'Numbering', 'Formatting']:
        doc.add_paragraph(section, style='List Bullet')
    
    add_heading_with_style(doc, 'Приложение Б. Логирование и отладка', 0)
    doc.add_paragraph('Система поддерживает подробное логирование всех операций.')
    doc.add_paragraph('Уровни логирования:')
    for level in ['DEBUG', 'INFO', 'WARNING', 'ERROR']:
        doc.add_paragraph(level, style='List Bullet')
    
    add_heading_with_style(doc, 'Приложение В. Примеры кода', 0)
    code_example = doc.add_paragraph('from doc_editor.editor import DocumentEditor\nfrom docx import Document')
    code_example.style = 'Normal'
    for run in code_example.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
    
    doc.save(output_path)
    print(f"✅ Полный документ создан: {output_path}")


def main():
    """Главная функция скрипта."""
    
    # Получить путь вывода
    if len(sys.argv) > 1 and not sys.argv[1].startswith('--'):
        output_path = sys.argv[1]
    else:
        output_path = 'doc_editor/tests/test_data/demo_document.docx'
    
    # Проверить флаги
    is_minimal = '--minimal' in sys.argv
    
    # Создать директорию если её нет
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    
    # Создать документ
    if is_minimal:
        create_minimal_demo_document(output_path)
    else:
        create_full_demo_document(output_path)
    
    print(f"\n📌 Путь к файлу: {Path(output_path).absolute()}")
    print(f"📌 Используйте этот файл для локального тестирования модуля")
    
    return 0


if __name__ == '__main__':
    sys.exit(main())
